"""
.docx exporter.

Converts extracted page content into an editable Word document with spell errors
marked up (red underline + yellow highlight). This is the translator-facing output:
open in Word or LibreOffice, edit inline, add translation notes.

Future: two-column layout (Tibetan | translation space), definition sidebars.
"""
import io
import logging
import re
from typing import Any

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX

logger = logging.getLogger(__name__)


def build_docx(
    pages: list[Any],              # list[PageContent]
    errors_by_page: dict[int, list[str]],  # page_number → list of error words
    filename: str = "document",
) -> bytes:
    """
    Build a .docx from extracted page content.

    Each page becomes a section. Words matching error_words are marked with
    a red underline and yellow highlight so they stand out for the translator.

    Returns the .docx as bytes.
    """
    doc = Document()

    # Document title
    doc.add_heading(f"OCR Output — {filename}", level=1)

    for page_content in pages:
        error_words = set(errors_by_page.get(page_content.page_number, []))

        # Page heading
        doc.add_heading(f"Page {page_content.page_number}", level=2)

        if not page_content.text.strip():
            doc.add_paragraph("[No text detected on this page]")
            continue

        # Process the page text line by line, then word by word
        for line in page_content.text.splitlines():
            if not line.strip():
                doc.add_paragraph("")
                continue

            para = doc.add_paragraph()
            _add_line_with_errors(para, line, error_words)

        # Page break between pages (except after the last page)
        if page_content.page_number < len(pages):
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _add_line_with_errors(para: Any, line: str, error_words: set[str]) -> None:
    """
    Split a line into individual syllables and delimiters, adding each as its
    own run. Only the specific syllables matching an error get error styling.

    Splitting on tshegs (་, U+0F0B) and spaces gives syllable-level granularity.
    Delimiters are emitted as unstyled runs so the text renders identically.
    """
    # Split on tsheg or space, keeping the delimiter as its own element.
    parts = re.split(r"(་| )", line)

    for part in parts:
        if not part:
            continue

        run = para.add_run(part)
        run.font.size = Pt(14)  # Tibetan script needs a larger size to be readable

        # Delimiters (tshegs, spaces) are never error syllables
        if part in ("་", " "):
            continue

        if part in error_words:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # dark red
            run.font.underline = True
            _set_highlight(run, WD_COLOR_INDEX.YELLOW)


def _set_highlight(run: Any, color: Any) -> None:
    """Apply a highlight color to a run via direct XML (python-docx doesn't expose this cleanly)."""
    rPr = run._r.get_or_add_rPr()
    highlight = OxmlElement("w:highlight")
    highlight.set(qn("w:val"), "yellow")
    rPr.append(highlight)
