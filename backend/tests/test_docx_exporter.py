"""
Tests for the DOCX exporter — specifically that error markup is applied at
syllable granularity rather than flagging an entire line.
"""
import pytest
from docx import Document
from docx.oxml.ns import qn

from app.pdf.docx_exporter import _add_line_with_errors, build_docx
from app.pdf.extractor import PageContent


def _run_texts_and_errors(para):
    """Return list of (text, is_error) for each non-empty run in a paragraph."""
    results = []
    for run in para.runs:
        if not run.text:
            continue
        # A run is marked as an error if it has a red font colour set
        is_error = run.font.color.rgb is not None and str(run.font.color.rgb) == "CC0000"
        results.append((run.text, is_error))
    return results


class TestAddLineWithErrors:

    def _make_para(self):
        doc = Document()
        return doc.add_paragraph()

    def test_single_error_syllable_only_that_syllable_marked(self):
        """Only the specific error syllable should be red; siblings must not be."""
        para = self._make_para()
        # Line with three syllables; middle one is the error
        _add_line_with_errors(para, "བོད་གཀར་ཡིག", {"གཀར"})
        runs = _run_texts_and_errors(para)

        error_runs = [t for t, err in runs if err]
        non_error_runs = [t for t, err in runs if not err]

        assert error_runs == ["གཀར"], f"Expected only གཀར flagged, got: {error_runs}"
        assert "བོད" in non_error_runs
        assert "ཡིག" in non_error_runs

    def test_no_errors_means_no_red_runs(self):
        """A clean line should have zero error-styled runs."""
        para = self._make_para()
        _add_line_with_errors(para, "བོད་ཡིག་གི་སྐད་ཡིག", set())
        runs = _run_texts_and_errors(para)
        error_runs = [t for t, err in runs if err]
        assert error_runs == []

    def test_whole_line_as_one_token_still_marks_only_error_syllable(self):
        """
        OCR text often has no spaces — the whole line is one tsheg-delimited
        string. Even so, only the matching syllable should be flagged.
        """
        para = self._make_para()
        # No spaces — mimics typical OCR output
        line = "རིན་པོ་ཆེར་གཀར་གྱི་མཐོང་"
        _add_line_with_errors(para, line, {"གཀར"})
        runs = _run_texts_and_errors(para)

        error_runs = [t for t, err in runs if err]
        non_error_content = [t for t, err in runs if not err and t not in ("་", " ")]

        assert error_runs == ["གཀར"], f"Expected only གཀར flagged, got: {error_runs}"
        # Every other syllable should be clean
        for syl in ["རིན", "པོ", "ཆེར", "གྱི", "མཐོང"]:
            assert syl in non_error_content, f"{syl} should be unmarked"

    def test_multiple_errors_each_marked_individually(self):
        para = self._make_para()
        _add_line_with_errors(para, "བོད་གཀར་ཡིག་གཀར་", {"གཀར"})
        runs = _run_texts_and_errors(para)
        error_runs = [t for t, err in runs if err]
        assert error_runs == ["གཀར", "གཀར"]

    def test_tshegs_and_spaces_never_marked_as_errors(self):
        para = self._make_para()
        _add_line_with_errors(para, "བོད་གཀར་ཡིག", {"གཀར"})
        runs = _run_texts_and_errors(para)
        for text, is_error in runs:
            if text in ("་", " "):
                assert not is_error, "Delimiters should never be styled as errors"

    def test_text_content_is_preserved(self):
        """Full text reconstructed from runs must equal the original line."""
        para = self._make_para()
        line = "རིན་པོ་ཆེར་གཀར་གྱི་"
        _add_line_with_errors(para, line, {"གཀར"})
        reconstructed = "".join(run.text for run in para.runs)
        assert reconstructed == line


class TestBuildDocx:

    def _pages_and_errors(self):
        page = PageContent(
            page_number=1,
            text="བོད་གཀར་ཡིག",
            is_scanned=False,
        )
        errors_by_page = {1: ["གཀར"]}
        return [page], errors_by_page

    def test_build_docx_returns_bytes(self):
        pages, errors = self._pages_and_errors()
        result = build_docx(pages, errors, "test.pdf")
        assert isinstance(result, bytes)
        assert len(result) > 0

    def test_build_docx_valid_docx_format(self):
        """Result should be openable as a valid docx."""
        import io
        pages, errors = self._pages_and_errors()
        result = build_docx(pages, errors, "test.pdf")
        doc = Document(io.BytesIO(result))
        # Should have at least a title paragraph and one content paragraph
        assert len(doc.paragraphs) >= 2
